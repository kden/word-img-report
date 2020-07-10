#!/usr/bin/env python

import os
import re
from glob import iglob
from docx import Document
from docx.shared import Inches

BASE_DIR = r'/home/caden/Downloads/onramp_result/'


def run_command(command):
    print(command)
    return os.system(command + ">/dev/null 2>&1")


def normalize_file_path(pathname):
    path, filename = os.path.split(pathname)
    filename = filename.lower().replace('-', '_')
    base = filename[0:-4]
    extension = filename[-3:]
    base = base.replace('.', '_')
    new_pathname = '/'.join([path, '.'.join([base, extension])])
    return new_pathname


# Create Word file

id_pattern_string = r'<dc:Identifier[^>]*>([0-9][^<]+)<'
id_pattern = re.compile(id_pattern_string)
title_pattern_string = r'<dc:Title[^>]*>([^<]*)<'
title_pattern = re.compile(title_pattern_string)
# <img src="images/cover.jpg" alt="Illustration" xml:space="preserve" id="img_00000" />
img_pattern_string_1 = r'<img[^>]* src=["\']([^"\']*)["\'][^>]*alt=["\']([^"\']*)["\'][^>]*>'
img_pattern_1 = re.compile(img_pattern_string_1)
img_pattern_string_2 = r'<img[^>]* alt=["\']([^"\']*)["\'][^>]*src=["\']([^"\']*)["\'][^>]*>'
img_pattern_2 = re.compile(img_pattern_string_2)

# For each DAISY file, unzip file
for zippath in iglob(BASE_DIR + '/*.zip'):
    path, filename = os.path.split(zippath)

    document = Document()
    document.add_heading('Bookshare + Math Detective Alt Text Test', 1)
    document.add_heading(filename, 2)

    basename = zippath.replace('.zip', '')
    run_command('rm -rf ' + basename)
    run_command('mkdir ' + basename)
    run_command('unzip -d ' + basename + ' ' + zippath)

    print(basename + '/images/*.*')
    for img_path in iglob(basename + '/images/*.*'):
        os.rename(img_path, normalize_file_path(img_path))

    for opfpath in iglob(basename + '/*.opf'):
        print(opfpath + '\n')
        # Get DAISY file metadata
        with open(opfpath, 'r') as file:
            file_data = file.read()
            identifiers = re.findall(id_pattern, file_data)
            for identifier in identifiers:
                print(identifier + '\n')
                isbn = identifier
            titles = re.findall(title_pattern, file_data)
            title = titles[0]
            print(title + '\n')

    # Write section header for DAISY file to Word document
    document.add_heading(isbn + ": " + title, 2)

    # Open DAISY XML file
    for xmlpath in iglob(basename + '/*.xml'):
        print(xmlpath + '\n')
        with open(xmlpath, 'r') as file:
            file_data = file.read()
            # Get list of img elements
            all_image_matches = []
            image_matches = re.finditer(img_pattern_1, file_data)
            for match in image_matches:
                all_image_matches.append((match.group(0), match.group(1), match.group(2)))
            image_matches = re.finditer(img_pattern_2, file_data)
            for match in image_matches:
                all_image_matches.append((match.group(0), match.group(2), match.group(1)))
            document.add_paragraph("Total number of image elements: " + str(len(all_image_matches)))
            table = document.add_table(rows=1, cols=3)
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Image'
            hdr_cells[1].text = 'Alt Text'
            hdr_cells[2].text = 'Full Tag'
            for match in all_image_matches:
                img_element = match[0]
                img_src =  match[1]
                img_path = normalize_file_path(basename + "/" + img_src)
                img_alt = match[2]
                row_cells = table.add_row().cells
                run = row_cells[0].add_paragraph().add_run()
                try:
                    run.add_picture(img_path, Inches(1.5), None)
                except FileNotFoundError:
                    print("File not found: " + basename + "/" + img_src)
                row_cells[1].text = img_alt
                row_cells[2].text = img_element
    output_name = filename.replace(".zip", ".docx")
    document.save(output_name)
